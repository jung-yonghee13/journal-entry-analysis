# -*- coding: utf-8 -*-
"""분개장 분석 파이프라인 (journal-entry-audit 스킬의 결정적 실행부).

사용법:
    python analyze_journal.py --input <분개장.xlsx|csv> [--outdir <출력폴더>]
                              [--sheet <시트명>] [--fy <회계연도>]

동작: 헤더 자동 탐지 → 컬럼 표준 매핑 → 이상 전표 탐지(A~L) → 계정별 추세 차트+해석
     → 현금흐름·차입금 분석 → 재무비율 → 원본 엑셀에 '이상전표' 시트 추가(백업 후)
     → 분석 보고서(md/docx/html) 생성. PDF는 Edge headless가 있으면 자동 변환.
점검 항목·문턱값 상세: ../references/checks.md
컬럼 매핑 규칙 상세:   ../references/column-mapping.md
"""
import argparse
import base64
import glob
import os
import re
import shutil
import subprocess
import sys
from datetime import date

import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import font_manager

# 콘솔이 cp949 등 비UTF-8이어도 출력이 죽지 않게
for _stream in (sys.stdout, sys.stderr):
    try:
        _stream.reconfigure(encoding="utf-8", errors="replace")
    except Exception:  # noqa: BLE001
        pass

# 한글 폰트: OS별 설치 폰트 중 첫 번째 사용 (없으면 경고 후 기본 폰트)
_KR_FONTS = ["Malgun Gothic", "AppleGothic", "Apple SD Gothic Neo",
             "NanumGothic", "Noto Sans CJK KR", "Noto Sans KR"]
_avail = {f.name for f in font_manager.fontManager.ttflist}
_font = next((f for f in _KR_FONTS if f in _avail), None)
if _font:
    plt.rcParams["font.family"] = _font
else:
    print("[경고] 한글 폰트 미발견 — 차트의 한글이 깨질 수 있습니다 (나눔고딕 등 설치 권장)")
plt.rcParams["axes.unicode_minus"] = False

# ════════════════════════════════════════════════════════════════
# 0. 컬럼 표준 매핑 (references/column-mapping.md 참조)
# ════════════════════════════════════════════════════════════════
SYNONYMS = {  # 표준필드: [동의어; 앞에 있을수록 우선]
    "회계일":     ["회계일자", "회계일", "전표일자", "기표일자", "전기일", "증빙일", "발생일",
                  "거래일자", "일자", "posting date", "gl date", "date"],
    "입력일시":   ["입력일시", "입력일", "등록일시", "생성일시", "entry date", "created"],
    "회계전표번호": ["회계전표번호", "전표번호", "전표no", "문서번호", "voucher no", "doc no"],
    "전표행번호":  ["전표행번호", "행번호", "라인번호", "line no"],
    "계정과목":   ["계정과목명", "계정과목", "계정명", "g/l account", "account"],
    "계정코드":   ["계정코드", "계정번호", "account code"],
    "차변금액":   ["차변금액", "차변액", "차변", "debit", "dr"],
    "대변금액":   ["대변금액", "대변액", "대변", "credit", "cr"],
    "금액":      ["전표금액", "금액", "amount"],
    "차대구분":   ["차대구분", "차대", "d/c", "dc구분"],
    "적요":      ["적요내용", "마스터적요", "적요", "description", "text"],
    "증빙":      ["증빙유형", "증빙구분", "증빙", "evidence"],
    "기표자":     ["기표자", "작성자", "입력자", "user"],
    "승인자":     ["승인자", "결재자", "approver"],
    "기표부서":   ["기표부서", "작성부서", "입력부서"],
    "승인부서":   ["승인부서", "결재부서"],
}
REQUIRED = ["회계일", "계정과목", "차변금액", "대변금액", "적요"]
HEADER_TOKENS = [s for v in SYNONYMS.values() for s in v]


def norm(x):
    return re.sub(r"[\s_()]", "", str(x)).lower()


def detect_header_row(path, sheet):
    """상위 8행 중 헤더 키워드가 가장 많이 등장하는 행을 헤더로 선택."""
    if path.lower().endswith((".xlsx", ".xls")):
        raw = pd.read_excel(path, sheet_name=sheet, header=None, nrows=8)
    else:
        raw = _read_csv(path, header=None, nrows=8)
    best, hits_best = 0, -1
    for i in range(len(raw)):
        cells = [norm(c) for c in raw.iloc[i].dropna()]
        hits = sum(any(norm(t) == c or norm(t) in c for t in HEADER_TOKENS) for c in cells)
        if hits > hits_best:
            best, hits_best = i, hits
    return best


def _read_csv(path, **kw):
    for enc in ("utf-8-sig", "cp949", "euc-kr", "utf-8"):
        try:
            return pd.read_csv(path, encoding=enc, **kw)
        except (UnicodeDecodeError, UnicodeError):
            continue
    return pd.read_csv(path, encoding="utf-8", errors="replace", **kw)


def map_columns(df):
    """동의어 사전으로 원본 컬럼 → 표준 필드 rename. 반환: (df, 매핑dict, 미해결 필수필드)"""
    cols = {c: norm(c) for c in df.columns}
    mapping, used = {}, set()
    EXACT_ONLY = {"금액", "차대구분"}  # 부분 일치 시 외화금액 등 오매핑 위험
    for std, syns in SYNONYMS.items():
        # 1차: 정규화 완전 일치 → 2차: 동의어가 컬럼명에 포함
        for exact in ((True,) if std in EXACT_ONLY else (True, False)):
            if std in mapping:
                break
            for syn in syns:
                s = norm(syn)
                for c, cn in cols.items():
                    if c in used:
                        continue
                    if (exact and cn == s) or (not exact and len(s) >= 2 and s in cn):
                        mapping[std] = c
                        used.add(c)
                        break
                if std in mapping:
                    break
    df = df.rename(columns={v: k for k, v in mapping.items()})
    # 금액 단일컬럼 + 차대구분 구조 → 차변/대변 분리
    if "차변금액" not in df.columns and {"금액", "차대구분"} <= set(df.columns):
        is_dr = df["차대구분"].astype(str).str.strip().str.upper().str[0].isin(list("차D1"))
        df["차변금액"] = df["금액"].where(is_dr, 0)
        df["대변금액"] = df["금액"].where(~is_dr, 0)
        mapping["차변금액"] = "금액+차대구분(분리)"
        mapping["대변금액"] = "금액+차대구분(분리)"
    missing = [f for f in REQUIRED if f not in df.columns]
    return df, mapping, missing


def _date_rate(s):
    """변환 성공률 — 연도가 타당(1990~2100)한 값만 성공으로 인정.
    (정수를 epoch 나노초로 오해석해 1970년이 되는 '가짜 성공'을 걸러낸다)"""
    return float((s.notna() & s.dt.year.between(1990, 2100)).mean())


def parse_dates(series):
    """문자열·YYYYMMDD 정수·Excel 직렬값을 모두 시도해 타당 성공률 최고 결과 반환."""
    best = pd.to_datetime(series, errors="coerce")
    best_r = _date_rate(best)
    if best_r < 0.8 and pd.api.types.is_numeric_dtype(series):
        for cand in (
            pd.to_datetime(series.astype("Int64").astype(str), format="%Y%m%d", errors="coerce"),
            pd.to_datetime(series, unit="D", origin="1899-12-30", errors="coerce"),
        ):
            r = _date_rate(cand)
            if r > best_r:
                best, best_r = cand, r
    return best.where(best.dt.year.between(1990, 2100))


def resolve_date_column(df):
    """매핑된 회계일의 변환 성공률이 낮으면 내용 기반으로 날짜 컬럼 재탐색."""
    if "회계일" in df.columns:
        parsed = parse_dates(df["회계일"])
        if _date_rate(parsed) >= 0.8:
            return parsed, "회계일(매핑)"
    best, best_rate, best_name = None, 0.0, None
    for c in df.columns:
        if c in ("차변금액", "대변금액", "금액"):
            continue
        p = parse_dates(df[c])
        rate = _date_rate(p)
        if rate > best_rate:
            best, best_rate, best_name = p, rate, c
    if best_rate >= 0.8:
        return best, f"{best_name}(내용 기반 재탐색)"
    raise SystemExit("[중단] 날짜 컬럼을 찾지 못했습니다. --sheet 또는 컬럼명을 확인하고, "
                     "사용자에게 날짜 컬럼을 확인받으세요. 컬럼 목록: " + ", ".join(map(str, df.columns)))


# ════════════════════════════════════════════════════════════════
# 1. 입력·전처리
# ════════════════════════════════════════════════════════════════
def load(path, sheet, fy_arg):
    is_excel = path.lower().endswith((".xlsx", ".xls"))
    if is_excel and sheet is None:
        xl = pd.ExcelFile(path)
        sheet = next((s for s in xl.sheet_names if "분개" in s or "전표" in s), xl.sheet_names[0])
    hdr = detect_header_row(path, sheet)
    df = (pd.read_excel(path, sheet_name=sheet, header=hdr) if is_excel
          else _read_csv(path, header=hdr))
    df, mapping, missing = map_columns(df)
    if missing and not (set(missing) <= {"회계일"}):
        raise SystemExit(f"[중단] 필수 컬럼 매핑 실패: {missing}. 사용자에게 컬럼 매핑을 확인받으세요. "
                         "컬럼 목록: " + ", ".join(map(str, df.columns)))
    df["회계일"], date_src = resolve_date_column(df)
    for c in ("차변금액", "대변금액"):
        df[c] = pd.to_numeric(df[c].astype(str).str.replace(r"[,\s₩]", "", regex=True),
                              errors="coerce").fillna(0)
    # 전표번호 보정
    if "회계전표번호" not in df.columns and "전표행번호" in df.columns:
        df["회계전표번호"] = df["전표행번호"].astype(str).str.replace(r"-\d+$", "", regex=True)
    if "회계전표번호" not in df.columns:
        raise SystemExit("[중단] 전표번호 컬럼을 찾지 못했습니다. 사용자에게 확인받으세요.")
    if "전표행번호" not in df.columns:
        df["전표행번호"] = (df["회계전표번호"].astype(str) + "-"
                        + (df.groupby("회계전표번호").cumcount() + 1).astype(str).str.zfill(3))
    df = df[df["계정과목"].notna() & (df["회계일"].notna() | (df["차변금액"] + df["대변금액"] > 0))]
    df = df.reset_index(drop=True)

    fy = fy_arg or int(df["회계일"].dt.year.mode()[0])
    df["연월"] = df["회계일"].dt.to_period("M").astype(str)
    df["요일"] = df["회계일"].dt.dayofweek
    try:
        import holidays
        kr = holidays.KR(years=sorted(df["회계일"].dt.year.dropna().unique().astype(int)))
        df["공휴일명"] = df["회계일"].dt.date.map(lambda d: kr.get(d, "") if pd.notna(d) else "")
    except ImportError:
        df["공휴일명"] = ""
        print("[경고] holidays 패키지 없음 — 공휴일 점검 생략(설치 권장: pip install holidays)")
    df["적요s"] = df["적요"].fillna("").astype(str).str.strip()
    df["금액"] = df[["차변금액", "대변금액"]].max(axis=1)

    print(f"[매핑] " + ", ".join(f"{k}←{v}" for k, v in mapping.items()))
    print(f"[매핑] 날짜: {date_src} / 시트: {sheet} / 헤더행: {hdr + 1}행째 / 회계연도(FY): {fy}")
    return df, fy, sheet


# ════════════════════════════════════════════════════════════════
# 2. 이상 전표 탐지 A~L (문턱값·근거: references/checks.md)
# ════════════════════════════════════════════════════════════════
SHORT = {
    "적요이상": "적요 공란/무의미", "기간외일자": "회계기간 외 일자", "날짜이상": "날짜 형식 오류",
    "주말분개": "주말 입력", "공휴일분개": "공휴일 입력", "결산전표이상": "월할계상 이탈/중복",
    "취소분개": "취소·역분개", "역분개의심": "당일 반대전표 쌍", "차대불일치": "차대변 불일치",
    "전표번호중복": "전표번호 중복", "자기승인": "셀프 승인", "라운드금액": "라운드 거액",
    "적요계정불일치": "적요-계정 불일치", "중복전표": "근접일 중복 전표",
    "가지급금미정산": "가지급금 미정산", "증빙누락": "거액 증빙 누락", "비경상대체": "비경상 계정 대체",
    "리스의심": "리스 회계처리 검토",
}
MISMATCH_RULES = [
    ("복리후생비", r"자택|사택|개인|가족|인테리어", "업무무관(개인성) 지출을 복리후생비 처리 의심"),
    (r"소모품비|사무용품비|잡비", r"골프|접대|답례|선물|회원권|경조|유흥", "접대성 지출의 일반 경비 분류 의심(기업업무추진비 검토)"),
    (r"수선비|수선유지비", r"증설|신설|신규|설치.?공사|공사비|리모델링", "자본적지출의 수선비(비용) 처리 의심(자산 계상 검토)"),
    (r"지급수수료", r"식대|회식|경조", "적요와 계정과목 성격 불일치"),
]


def detect(df, fy):
    fy_end = pd.Timestamp(f"{fy}-12-31")
    issues = []

    def flag(idx, kind, desc):
        issues.append({"idx": idx, "이상유형": kind, "이상징후": SHORT[kind], "문제설명": desc})

    def flag_voucher(vno, kind, desc):
        for idx in df[df["회계전표번호"] == vno].index:
            flag(idx, kind, desc)

    # A. 적요 이상
    for idx, r in df.iterrows():
        s = r["적요s"]
        if not s:
            flag(idx, "적요이상", "적요가 비어 있음")
        elif s.lower() in ("test", "테스트", "asdf", ".") or len(s) < 3:
            flag(idx, "적요이상", f"무의미한 적요: '{s}'")
    # B. 날짜
    for idx, r in df[df["회계일"].dt.year != fy].iterrows():
        flag(idx, "기간외일자", f"회계연도({fy}) 밖의 일자 {r['회계일'].date()} — 기간 귀속 확인 필요")
    for idx in df[df["회계일"].isna()].index:
        flag(idx, "날짜이상", "회계일 변환 실패(형식 오류)")
    # C. 주말·공휴일
    for idx, r in df.iterrows():
        if r["요일"] >= 5:
            flag(idx, "주말분개", f"{r['회계일'].date()} {'토' if r['요일'] == 5 else '일'}요일 분개")
        if r["공휴일명"]:
            flag(idx, "공휴일분개", f"{r['회계일'].date()} 공휴일({r['공휴일명']}) 분개")
    # D. 결산전표 (월할 정책 인식 → 이탈·중복만)
    closing = df[df["계정과목"].astype(str).str.contains("감가상각|대손상각|대손충당|퇴직급여|법인세", na=False)]
    policy_notes = []
    for name, g in closing.groupby("계정과목"):
        if g["연월"].nunique() >= 6:
            policy_notes.append(f"{name}: 매월 계상({g['연월'].nunique()}개월) — 월할 정책으로 판단, 규칙 발생분 제외")
            for idx, r in g[~g["회계일"].dt.is_month_end].iterrows():
                flag(idx, "결산전표이상", f"월할 계상 계정({name})이 말일이 아닌 {r['회계일'].date()}에 계상")
            dup = g["연월"].value_counts()
            for ym in dup[dup > 1].index:
                for idx in g[g["연월"] == ym].index:
                    flag(idx, "결산전표이상", f"월할 계상 계정({name})이 {ym}에 {dup[ym]}회 계상 — 이중 계상 또는 처분 여부 확인")
        else:
            for idx, r in g[g["회계일"] != fy_end].iterrows():
                flag(idx, "결산전표이상", f"결산 성격 계정({name})을 기말 외 일자({r['회계일'].date()})에 계상")
    # E. 취소·역분개
    for vno in df.loc[df["적요s"].str.contains("취소|역분개|반품|정정"), "회계전표번호"].unique():
        g = df[df["회계전표번호"] == vno]
        writer = g["기표자"].iloc[0] if "기표자" in df.columns else "?"
        d = f"취소/역분개 전표 (기표자 {writer})"
        if g["회계일"].iloc[0] == fy_end:
            d += " — 기말일 대형 취소, 이익조정 위험"
        flag_voucher(vno, "취소분개", d)
    d_ = df[df["차변금액"] > 0][["회계전표번호", "회계일", "계정과목", "차변금액"]].rename(columns={"차변금액": "금액"})
    c_ = df[df["대변금액"] > 0][["회계전표번호", "회계일", "계정과목", "대변금액"]].rename(columns={"대변금액": "금액"})
    sd = d_.merge(c_, on=["회계일", "계정과목", "금액"], suffixes=("_D", "_C"))
    for _, r in sd[sd["회계전표번호_D"] != sd["회계전표번호_C"]].iterrows():
        for vno in (r["회계전표번호_D"], r["회계전표번호_C"]):
            for idx in df[(df["회계전표번호"] == vno) & (df["계정과목"] == r["계정과목"])].index:
                if "취소" not in df.loc[idx, "적요s"]:
                    flag(idx, "역분개의심", f"같은 날 {r['계정과목']} {r['금액']:,.0f}원 반대 방향 쌍"
                                       f"({r['회계전표번호_D']}↔{r['회계전표번호_C']})")
    # F. 정합성
    vsum = df.groupby("회계전표번호")[["차변금액", "대변금액"]].sum()
    for vno, r in vsum[(vsum["차변금액"] - vsum["대변금액"]).abs() > 0].iterrows():
        flag_voucher(vno, "차대불일치", f"차변합 {r['차변금액']:,.0f} ≠ 대변합 {r['대변금액']:,.0f}"
                                    f" (차이 {r['차변금액'] - r['대변금액']:,.0f})")
    for idx in df[df.duplicated("전표행번호", keep=False)].index:
        flag(idx, "전표번호중복", f"전표행번호 중복: {df.loc[idx, '전표행번호']}")
    if {"기표자", "승인자"} <= set(df.columns):
        for idx, r in df[df["기표자"].astype(str) == df["승인자"].astype(str)].iterrows():
            flag(idx, "자기승인", f"기표자와 승인자가 동일({r['기표자']})")
    for idx in df[(df["금액"] >= 50_000_000) & (df["금액"] % 10_000_000 == 0)].index:
        d = f"라운드 금액 {df.loc[idx, '금액']:,.0f}원 — 실재성·근거 확인 대상"
        if df.loc[idx, "회계일"] >= fy_end - pd.Timedelta(days=7):
            d += " (기말 직전 집중 계상)"
        flag(idx, "라운드금액", d)
    # G. 적요-계정 불일치
    for acc_pat, memo_pat, why in MISMATCH_RULES:
        for idx, r in df[df["계정과목"].astype(str).str.contains(acc_pat)
                         & df["적요s"].str.contains(memo_pat)].iterrows():
            flag_voucher(r["회계전표번호"], "적요계정불일치", f"'{r['적요s']}' → {why}")
    # H. 근접일 중복 전표
    for (acc, amt, memo), g in df[df["차변금액"] >= 1_000_000].groupby(["계정과목", "차변금액", "적요s"]):
        if not memo or g["회계전표번호"].nunique() < 2:
            continue
        g = g.sort_values("회계일")
        if (g["회계일"].diff().dt.days.dropna() <= 7).any():
            vnos = g["회계전표번호"].unique()
            for vno in vnos:
                flag_voucher(vno, "중복전표", f"동일 계정·금액·적요({acc}, {amt:,.0f}, '{memo}') "
                                          f"전표가 7일 이내 반복({', '.join(vnos)}) — 이중 계상 의심")
    # I. 가지급금 미정산
    for acc in ("가지급금", "가불금", "전도금"):
        if acc in df["계정과목"].values:
            g = df[df["계정과목"] == acc]
            if g["차변금액"].sum() > 0 and g["대변금액"].sum() == 0:
                for _, r in g[g["차변금액"] > 0].iterrows():
                    flag_voucher(r["회계전표번호"], "가지급금미정산",
                                 f"{acc} {r['차변금액']:,.0f}원 발생 후 연중 정산(대변) 전표 없음 — "
                                 "업무무관 가지급금·인정이자 검토")
    # J. 거액 증빙 누락
    if "증빙" in df.columns:
        for idx, r in df[df["증빙"].isna() & (df["금액"] >= 5_000_000)].iterrows():
            d = f"증빙란 공란 + 거액({r['금액']:,.0f}원)"
            if re.search("기업업무추진비|접대비|기부금|지급수수료|여비교통비", str(r["계정과목"])):
                d += f" — 민감 계정({r['계정과목']}): 한도초과·가공경비 검토"
            flag_voucher(r["회계전표번호"], "증빙누락", d)
    # K. 비경상 대체 (부채 차변 ↔ 수익 대변)
    for vno, g in df.groupby("회계전표번호"):
        liab = g[(g["차변금액"] > 0) & g["계정과목"].astype(str).str.contains("외상매입금|미지급|차입금|예수금|선수금")]
        rev = g[(g["대변금액"] > 0) & g["계정과목"].astype(str).str.contains("잡이익|잡수익|채무면제")]
        if len(liab) and len(rev):
            flag_voucher(vno, "비경상대체", f"부채({liab['계정과목'].iloc[0]}) {liab['차변금액'].sum():,.0f}원을 "
                                        f"수익({rev['계정과목'].iloc[0]})으로 대체 — 근거 문서·채무면제 실질 확인 필요")
    # L. 리스 의심
    lt = df[df["계정과목"].astype(str).str.contains("지급수수료|임차료|렌탈|사용료|잡비") & (df["차변금액"] > 0)].copy()
    for idx, r in lt[lt["적요s"].str.contains(r"리스|렌탈|렌트|장기임대|월\s?사용료|약정|할부")].iterrows():
        flag_voucher(r["회계전표번호"], "리스의심",
                     f"'{r['적요s']}' — 리스·렌탈 성격 적요가 {r['계정과목']}로 비용 처리됨. 사용권자산·리스부채 계상 대상 여부 확인")
    lt["적요norm"] = lt["적요s"].str.replace(r"\d+월?|\(.*?\)", "", regex=True).str.strip()
    for memo, g in lt.groupby("적요norm"):
        if not memo or g["연월"].nunique() < 6:
            continue
        amts = g.groupby("연월")["차변금액"].sum()
        if amts.mean() > 0 and amts.std() / amts.mean() <= 0.10:
            acc = g["계정과목"].mode()[0]
            d = (f"'{memo}' 월 {amts.mean() / 1e6:,.1f}백만원 × {g['연월'].nunique()}개월 정기 고정 지급"
                 f"(연간 {g['차변금액'].sum() / 1e8:,.2f}억) — 고정 리스료 성격. K-IFRS 1116 적용 시 "
                 "사용권자산·리스부채 계상 대상 여부 확인" + (", 지급수수료 분류 오류 가능성" if "수수료" in acc else ""))
            for vno in g["회계전표번호"].unique():
                flag_voucher(vno, "리스의심", d)

    issues_df = pd.DataFrame(issues).drop_duplicates(subset=["idx", "이상유형", "문제설명"])
    return issues_df, policy_notes


# ════════════════════════════════════════════════════════════════
# 3. 계정별 추세 차트 + 해석 (급변 이중 기준·주기성·사유 추론)
# ════════════════════════════════════════════════════════════════
REASON_RULES = [
    (r"부가세.*(신고|정리)|기예정|기확정", "분기 부가세 신고(예정·확정) 정리에 따른 주기적 집계 — 정상적 세무 일정"),
    (r"단기차입|차입.*상환|운영자금", "차입금 유입·상환에 따른 일시 변동 — 약정서·이사회 승인 확인"),
    (r"설비대금|증설|신규.*공사", "설비 투자 지출 — 자본적지출의 자산 계상 여부 확인"),
    (r"상각누계.*제거|처분", "유형자산 처분에 따른 상각누계액 제거 — 처분손익 계산 확인"),
    (r"경영자문|컨설팅|자문료", "자문·컨설팅 수수료 일시 지급 — 용역 실재성·특수관계자 여부 확인 필요"),
    (r"급여|상여|원천세|보험", "급여·상여 및 원천세 납부 주기에 따른 변동"),
    (r"일괄", "기말·특정 시점 일괄 계상 — 기간 배분의 적정성 확인"),
    (r"매출.*취소|반품", "매출 취소·반품 — 취소 사유와 증빙 확인"),
    (r"어음수취", "대금 회수 수단(어음) 비중 변화"),
    (r"회수", "매출채권 회수 집중"),
    (r"매출", "매출 관련 대금 흐름 변동"),
    (r"매입", "매입 물량 변동"),
]


def infer_reason(memo):
    for pat, why in REASON_RULES:
        if re.search(pat, str(memo)):
            return why
    return None


def build_charts(df, issues_df, fy, chart_dir):
    months = [f"{fy}-{m:02d}" for m in range(1, 13)]
    dfy = df[df["회계일"].dt.year == fy]
    monthly = dfy.groupby(["계정과목", "연월"])[["차변금액", "대변금액"]].sum().reset_index()
    monthly["발생액"] = monthly[["차변금액", "대변금액"]].max(axis=1)
    acc_total = dfy.groupby("계정과목")["금액"].sum().sort_values(ascending=False)
    targets = acc_total[acc_total >= 100_000_000].head(30).index.tolist()

    iss = issues_df.copy()
    iss["계정과목"] = df.loc[iss["idx"], "계정과목"].values
    iss["연월"] = df.loc[iss["idx"], "연월"].values
    issue_ma = iss.groupby(["계정과목", "연월"])["이상징후"].agg(lambda s: ", ".join(sorted(set(s))))
    flagged = set(df.loc[issues_df["idx"].unique(), "회계전표번호"]) if len(issues_df) else set()

    charts = []
    for acc in targets:
        g_acc = dfy[dfy["계정과목"] == acc]
        s = monthly[monthly["계정과목"] == acc].set_index("연월")["발생액"].reindex(months, fill_value=0)
        base = s[s > 0].median() if (s > 0).any() else 0
        spikes = []
        for i, m in enumerate(months):
            why = []
            prev = s.iloc[i - 1] if i > 0 else None
            if prev and prev > 0 and abs(s[m] - prev) / prev >= 0.35 and abs(s[m] - prev) >= 20_000_000:
                why.append(f"전월 대비 {(s[m] - prev) / prev * 100:+.0f}%")
            if base > 0 and s[m] >= 1.7 * base and s[m] >= 30_000_000:
                why.append(f"월 중앙값의 {s[m] / base:.1f}배")
            if why:
                spikes.append((m, why))
        up = [(m, w) for m, w in spikes if base == 0 or s[m] >= base]
        up_months = {m for m, _ in up}

        fig, ax = plt.subplots(figsize=(7.2, 2.9))
        colors = ["#DD8452" if m in up_months else "#4C72B0" for m in months]
        ax.bar(range(12), s.values / 1e6, color=colors, width=0.62)
        ax.plot(range(12), s.values / 1e6, color="#C44E52", marker="o", ms=4, lw=1.2)
        ax.set_xticks(range(12)); ax.set_xticklabels([m[5:] + "월" for m in months], fontsize=8)
        ax.set_title(f"{acc} — 월별 발생액(백만원)", fontsize=11)
        ax.tick_params(labelsize=8); ax.grid(axis="y", alpha=.3)
        fig.tight_layout()
        fname = os.path.join(chart_dir, f"chart_{re.sub(r'[^가-힣A-Za-z0-9]', '_', acc)}.png")
        fig.savefig(fname, dpi=110); plt.close(fig)

        parts, active = [], int((s > 0).sum())
        up_m = sorted(int(m[5:]) for m in up_months)
        if len(up_m) >= 3 and set(up_m) <= {1, 4, 7, 10}:
            drv = g_acc[g_acc["연월"].isin([f"{fy}-{mm:02d}" for mm in up_m])].nlargest(1, "금액")
            reason = infer_reason(drv["적요s"].iloc[0]) if len(drv) else None
            parts.append(f"{'·'.join(map(str, up_m))}월 주기적 급증(예: {drv['적요s'].iloc[0]} "
                         f"{drv['금액'].iloc[0] / 1e8:,.1f}억)" + (f" → {reason}" if reason else ""))
        elif active <= 3:
            desc = []
            for m in [m for m in months if s[m] > 0]:
                drv = g_acc[g_acc["연월"] == m].nlargest(1, "금액")
                memo = drv["적요s"].iloc[0] if len(drv) else ""
                r = infer_reason(memo)
                desc.append(f"{int(m[5:])}월 {s[m] / 1e8:,.1f}억 '{memo}'" + (f" → {r}" if r else ""))
            parts.append(f"연중 {active}개월만 발생: " + " / ".join(desc))
        elif up:
            for m, why in sorted(up, key=lambda x: -(s[x[0]] - base))[:2]:
                drv = g_acc[g_acc["연월"] == m].nlargest(1, "금액")
                memo = drv["적요s"].iloc[0] if len(drv) else ""
                amt = drv["금액"].iloc[0] if len(drv) else 0
                vno = drv["회계전표번호"].iloc[0] if len(drv) else ""
                t = f"{int(m[5:])}월 급증({', '.join(why)}) — 주요 원인: '{memo}' {amt / 1e8:,.1f}억"
                reason = infer_reason(memo)
                if reason:
                    t += f" → {reason}"
                if vno in flagged:
                    t += " [이상전표로 탐지된 전표]"
                parts.append(t)
        else:
            cv = s[s > 0].std() / s[s > 0].mean() if active > 1 else 0
            h1, h2 = s[:6].mean(), s[6:].mean()
            if h1 > 0 and abs(h2 - h1) / h1 >= 0.3:
                parts.append(f"하반기 평균이 상반기 대비 {(h2 - h1) / h1 * 100:+.0f}% {'증가' if h2 > h1 else '감소'}")
            else:
                parts.append("급변 월 없음 — " + ("연중 안정적 발생" if cv < 0.35 else "완만한 월별 변동"))
            parts.append(f"최대 {int(s.idxmax()[5:])}월 {s.max() / 1e8:,.1f}억원")
        if acc in issue_ma.index.get_level_values(0):
            items = [f"{int(m[5:])}월({v})" for m, v in issue_ma.loc[acc].sort_index().items()]
            t = "⚠ 이상 전표: " + " / ".join(items[:6])
            if len(items) > 6:
                t += f" 외 {len(items) - 6}개월 (이상전표 시트 참조)"
            parts.append(t)
        charts.append((acc, fname, " · ".join(parts)))
    return charts, months


# ════════════════════════════════════════════════════════════════
# 4. 현금흐름 + 차입금 / 5. 재무비율
# ════════════════════════════════════════════════════════════════
def cashflow(df, fy, months, chart_dir):
    dfy = df[df["회계일"].dt.year == fy]
    CASH, FIN = r"현금$|보통예금|당좌예금", r"차입금|사채|자본금|배당"
    INV = r"기계장치|차량운반구|비품|건물|토지|구축물|대여금|투자|보증금|소프트웨어"
    recs = []
    for _, g in dfy.groupby("회계전표번호"):
        cash = g[g["계정과목"].astype(str).str.contains(CASH)]
        if cash.empty:
            continue
        delta = cash["차변금액"].sum() - cash["대변금액"].sum()
        others = g[~g["계정과목"].astype(str).str.contains(CASH)]
        if delta == 0:
            continue
        if others.empty:
            cat, top_acc, memo = "기타(현금간 대체)", "-", ""
        else:
            top = others.loc[others["금액"].idxmax()]
            top_acc, memo = str(top["계정과목"]), str(top["적요s"])
            cat = ("재무활동" if re.search(FIN, top_acc)
                   else "투자활동" if re.search(INV, top_acc) else "영업활동")
        recs.append({"연월": g["연월"].iloc[0], "활동": cat, "순현금": delta,
                     "상대계정": top_acc, "적요": memo, "절대값": abs(delta)})
    cf = pd.DataFrame(recs)
    piv = (cf.pivot_table(index="연월", columns="활동", values="순현금", aggfunc="sum")
           .reindex(months, fill_value=0).fillna(0))
    for c in ("영업활동", "투자활동", "재무활동"):
        if c not in piv.columns:
            piv[c] = 0
    piv = piv[["영업활동", "투자활동", "재무활동"]]
    piv["순증감"] = piv.sum(axis=1)
    piv["누적"] = piv["순증감"].cumsum()

    fig, ax = plt.subplots(figsize=(9.5, 3.6))
    for i, (col, c) in enumerate([("영업활동", "#4C72B0"), ("투자활동", "#DD8452"), ("재무활동", "#55A868")]):
        ax.bar([x + (i - 1) * 0.27 for x in range(12)], piv[col] / 1e8, width=0.27, label=col, color=c)
    ax2 = ax.twinx()
    ax2.plot(range(12), piv["누적"] / 1e8, color="#333", marker="o", ms=4, lw=1.6)
    ax2.set_ylabel("누적(억원)", fontsize=9); ax2.tick_params(labelsize=8)
    ax.axhline(0, color="#999", lw=.8)
    ax.set_xticks(range(12)); ax.set_xticklabels([m[5:] + "월" for m in months], fontsize=8)
    ax.set_ylabel("월별 순현금(억원)", fontsize=9); ax.tick_params(labelsize=8)
    ax.legend(fontsize=8, loc="upper left"); ax.grid(axis="y", alpha=.3)
    ax.set_title(f"활동별 월별 순현금흐름과 누적 추이 ({fy})", fontsize=11)
    fig.tight_layout()
    cf_chart = os.path.join(chart_dir, "현금흐름_월별.png")
    fig.savefig(cf_chart, dpi=110); plt.close(fig)

    loan = df[df["계정과목"].astype(str).str.contains("차입금|사채")].sort_values("회계일")
    loan_rows = []
    for _, r in loan.iterrows():
        kind = "차입" if r["대변금액"] > 0 else "상환"
        near = loan[(loan["회계일"] - r["회계일"]).abs().dt.days <= 7]
        note = ("근접일 반대 거래 존재 — 대환(차환) 가능성 확인"
                if ((kind == "차입" and (near["차변금액"] > 0).any())
                    or (kind == "상환" and (near["대변금액"] > 0).any())) else "")
        loan_rows.append((str(r["회계일"].date()), kind, f"{r['금액']:,.0f}", r["적요s"], note))

    interp = []
    op, inv, fin = piv["영업활동"].sum(), piv["투자활동"].sum(), piv["재무활동"].sum()
    interp.append(f"연간 순현금흐름: 영업 {op / 1e8:+,.1f}억 / 투자 {inv / 1e8:+,.1f}억 / 재무 {fin / 1e8:+,.1f}억 "
                  f"→ 합계 {(op + inv + fin) / 1e8:+,.1f}억. "
                  + ("영업활동에서 현금을 창출해 투자·재무 지출을 감당하는 구조"
                     if op > 0 and op >= -inv else
                     "영업 현금창출이 부족하거나 외부 자금에 의존하는 흐름 — 자금계획 점검 필요"
                     " (기초잔액 미포함 분개장 기준이므로 실제 자금 사정은 잔액 확인 필요)"))
    if len(cf):
        for m in piv["순증감"].abs().sort_values(ascending=False).head(2).index:
            mv = cf[cf["연월"] == m]
            if mv.empty:
                continue
            drv = mv.loc[mv["절대값"].idxmax()]
            interp.append(f"{int(m[5:])}월 순증감 {piv.loc[m, '순증감'] / 1e8:+,.1f}억 — 최대 요인: "
                          f"[{drv['활동']}] '{drv['적요']}' {drv['순현금'] / 1e8:+,.1f}억({drv['상대계정']})")
    if len(loan):
        차입, 상환 = loan["대변금액"].sum(), loan["차변금액"].sum()
        ev = " / ".join(f"{d[5:]} {k} {a}원" for d, k, a, _, _ in loan_rows[:4])
        interp.append(f"차입금: 연중 차입 {차입 / 1e8:,.1f}억, 상환 {상환 / 1e8:,.1f}억, 순증 {(차입 - 상환) / 1e8:+,.1f}억"
                      f" ({ev}) — 잔여 차입금의 만기·이자율 조건과 상환 계획 확인 권고")
    inv_ev = cf[cf["활동"] == "투자활동"].sort_values("절대값", ascending=False).head(3) if len(cf) else []
    if len(inv_ev):
        interp.append("주요 투자활동: " + " / ".join(
            f"{int(r['연월'][5:])}월 '{r['적요']}' {r['순현금'] / 1e8:+,.1f}억" for _, r in inv_ev.iterrows()))
    return piv, cf_chart, loan_rows, interp


def ratios(df, fy):
    dfy = df[df["회계일"].dt.year == fy]
    bal = dfy.groupby("계정과목")[["차변금액", "대변금액"]].sum()
    bal["잔액"] = bal["차변금액"] - bal["대변금액"]

    def s_(pat, sign=1):
        got = bal[bal.index.astype(str).str.contains(pat)]
        return sign * got["잔액"].sum()

    유동자산 = s_("현금$|예금|외상매출금|받을어음|미수금|선급|가지급금|부가세대급금|원재료|재고|상품|제품$|대손충당금")
    유동부채 = s_("외상매입금|미지급|예수금|부가세예수금|단기차입금|선수금", -1)
    매출채권 = s_("외상매출금|받을어음|대손충당금")
    매출액 = s_("매출$|제품매출|상품매출", -1)
    재고 = s_("재고|원재료|상품$|제품$|재공품")
    매출원가 = s_("매출원가")
    rows = []
    if 유동부채 > 0:
        v = 유동자산 / 유동부채 * 100
        c = "양호(200% 이상)" if v >= 200 else "보통(100~200%)" if v >= 100 else "단기 지급능력 주의(100% 미만)"
        rows.append(("유동비율", f"{v:,.1f}%", f"유동자산 {유동자산 / 1e8:,.1f}억 / 유동부채 {유동부채 / 1e8:,.1f}억 — {c}"))
    else:
        rows.append(("유동비율", "계산 불가", "유동부채 계정 잔액 없음"))
    if 매출채권 > 0 and 매출액 > 0:
        t = 매출액 / 매출채권
        days = 365 / t
        rows.append(("매출채권회전율", f"{t:,.1f}회", f"매출 {매출액 / 1e8:,.1f}억 / 채권 {매출채권 / 1e8:,.1f}억"))
        rows.append(("매출채권회수기간", f"{days:,.0f}일",
                     "회수기간이 김 — 채권 회수 관리 점검 권고" if days > 120 else "통상 범위"))
    else:
        rows.append(("매출채권회전율", "계산 불가", "매출 또는 매출채권 계정 잔액 없음"))
    if 재고 > 0 and 매출원가 > 0:
        rows.append(("재고자산회전율", f"{매출원가 / 재고:,.1f}회", f"매출원가 {매출원가 / 1e8:,.1f}억 / 재고 {재고 / 1e8:,.1f}억"))
    else:
        rows.append(("재고자산회전율", "계산 불가",
                     f"매출원가 계정 부재(결산 대체 미반영){' — 재고 순증 ' + format(재고 / 1e8, ',.1f') + '억' if 재고 else ''}"))
    return rows


# ════════════════════════════════════════════════════════════════
# 5b. M — AI 정황 스캔용 후보 추출 (결정적)
#     100만 행급 데이터에서도 AI가 원시 행을 읽지 않도록, 관점별 후보를
#     스크립트가 먼저 좁혀 파일로 저장한다. AI는 이 파일만 판독한다.
#     토큰 소비가 데이터 행 수가 아니라 후보 수(관점별 상한 N)에 비례한다.
# ════════════════════════════════════════════════════════════════
M_TOP_N = 50  # 관점별 표시 상한 (금액 상위) — 초과분은 건수만 기록

def m_candidates(df, issues_df, outdir):
    flagged = set(df.loc[issues_df["idx"].unique(), "회계전표번호"]) if len(issues_df) else set()
    L = ["# AI 부정 위험 정황 스캔 — 후보 목록 (스크립트 결정적 추출)", "",
         "> AI는 원시 분개장이 아니라 **이 후보 목록만** 판독해 정황 여부를 판단한다.",
         f"> 관점별 금액 상위 {M_TOP_N}건 상한. 기탐지=규칙(A~L)으로 이미 이상전표에 포함됨.", ""]
    counts = {}

    def emit(title, sub, extra_cols=()):
        sub = sub.drop_duplicates("회계전표번호")
        total = len(sub)
        counts[title] = total
        L.append(f"## {title} — 후보 {total}건" + (f" (금액 상위 {M_TOP_N}건만 표시, 나머지 {total - M_TOP_N}건 생략)" if total > M_TOP_N else ""))
        if total == 0:
            L.append("- 없음\n")
            return
        cols = ["회계전표번호", "회계일", "계정과목", "금액", "적요s"] + [c for c in extra_cols if c in sub.columns]
        show = sub.nlargest(min(total, M_TOP_N), "금액")
        L.append("| " + " | ".join(cols) + " | 기탐지 |")
        L.append("|" + "---|" * (len(cols) + 1))
        for _, r in show.iterrows():
            vals = [f"{r[c]:,.0f}" if c == "금액" else str(r[c].date()) if c == "회계일" else str(r[c]) for c in cols]
            L.append("| " + " | ".join(vals) + " | " + ("Y" if r["회계전표번호"] in flagged else "**N**") + " |")
        L.append("")

    # M1. 위법·부정 암시 적요
    emit("M1. 부정 암시 적요(리베이트·사례비·커미션 등)",
         df[df["적요s"].str.contains("리베이트|사례비|커미션|현금화|뒷돈|비자금|급행료|성공보수")], ("기표자",))
    # M2. 특수관계자 정황
    emit("M2. 특수관계자 정황(대표·임원·가족·개인성 지출)",
         df[df["적요s"].str.contains("대표이사|대표 |임원|사장|자택|사택|개인|가족")], ("기표자",))
    # M3. 권한·통제 우회 (부서 컬럼이 있을 때만)
    if {"기표자", "기표부서"} <= set(df.columns):
        usual = df.groupby("기표자")["기표부서"].agg(lambda s: s.mode()[0])
        odd = df[df["기표부서"] != df["기표자"].map(usual)]
        if {"승인자", "승인부서"} <= set(df.columns):
            usual_ap = df.groupby("승인자")["승인부서"].agg(lambda s: s.mode()[0])
            odd = pd.concat([odd, df[df["승인부서"] != df["승인자"].map(usual_ap)]])
        emit("M3. 권한·통제 우회(평소 부서와 다른 기표·승인)", odd, ("기표자", "기표부서", "승인자", "승인부서"))
    else:
        L.append("## M3. 권한·통제 우회 — 부서 컬럼 없음, 점검 생략\n")
        counts["M3(생략)"] = 0
    # M4. 승인 한도 근접 금액 반복 (450~500만 / 900~1,000만)
    amt = df["차변금액"]
    near = df[((amt >= 4_500_000) & (amt < 5_000_000)) | ((amt >= 9_000_000) & (amt < 10_000_000))]
    near = near[~near["계정과목"].astype(str).str.contains("매출|매입|예금|현금|원재료|부가세", na=False)]
    grp = (near.groupby(["계정과목"] + (["기표자"] if "기표자" in df.columns else []))
           .agg(건수=("차변금액", "size"), 합계=("차변금액", "sum"),
                일자=("회계일", lambda s: ", ".join(sorted(set(s.dt.strftime("%m/%d")))[:8]))))
    grp = grp[grp["건수"] >= 2].sort_values("합계", ascending=False).head(M_TOP_N)
    counts["M4. 한도 근접 반복"] = len(grp)
    L.append(f"## M4. 승인 한도 근접 금액 반복 — 그룹 {len(grp)}건")
    L.append("(근접일 연속 분할인지, 수개월 간격의 정상 발생인지는 AI가 일자 간격으로 판단)")
    if len(grp):
        g2 = grp.reset_index()
        L.append("| " + " | ".join(map(str, g2.columns)) + " |")
        L.append("|" + "---|" * len(g2.columns))
        for _, r in g2.iterrows():
            L.append("| " + " | ".join(f"{v:,.0f}" if isinstance(v, (int, float)) else str(v) for v in r) + " |")
    else:
        L.append("- 없음")
    L.append("")
    # M5. 저빈도 거래처의 거액 거래
    trader = df["적요s"].str.extract(r"^([가-힣A-Za-z0-9]+)")[0]
    freq = trader.map(df.groupby(trader)["회계전표번호"].nunique())
    rare = df[(freq <= 2) & (df["금액"] >= 30_000_000) & (df["차변금액"] > 0)
              & ~df["계정과목"].astype(str).str.contains("예금|현금|부가세|매출", na=False)]
    emit("M5. 저빈도 거래처(연 2회 이하)의 거액(3천만↑) 거래", rare, ("기표자",))

    path = os.path.join(outdir, "AI정황_후보.md")
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(L) + "\n")
    return path, counts


# ════════════════════════════════════════════════════════════════
# 6. 시트 추가 + 보고서(md/docx/html/pdf)
# ════════════════════════════════════════════════════════════════
RECOMMEND_MAP = {
    "취소분개": "취소·역분개 → 원거래 증빙 추적, 반품·취소의 실물 증빙과 수정세금계산서 확인. 기말일 취소는 이익조정 위험이므로 차기 초 재계상 여부 추적",
    "비경상대체": "비경상 계정 대체 → 채무면제 등의 법적 근거(합의서) 확인, 기표자·승인자 부서 적정성 점검",
    "적요계정불일치": "적요-계정 불일치 → 업무관련성·자본화 여부 재검토, 세무조정 영향 평가",
    "가지급금미정산": "가지급금 미정산 → 정산 요구, 업무무관 가지급금 인정이자 계산, 대여금 전환 여부 결정",
    "중복전표": "근접일 중복 전표 → 세금계산서 원본 대조로 이중 계상 확인, 중복 시 취소분개",
    "증빙누락": "거액 증빙 누락 → 증빙 보완 요구, 민감 계정은 한도초과·가공경비 검토",
    "자기승인": "자기승인 → 승인 절차 재설계(기표·승인 분리), 해당 전표 재승인",
    "차대불일치": "차대변 불일치 → 즉시 내부통제 담당자 보고 및 전표 수정",
    "주말분개": "주말·공휴일 입력 → 입력 경위(자동 배치 여부)와 시스템 접근 기록 확인",
    "공휴일분개": "주말·공휴일 입력 → 입력 경위(자동 배치 여부)와 시스템 접근 기록 확인",
    "적요이상": "적요 공란·무의미 → 기표자 소명, 증빙 대조 후 적요 보완",
    "기간외일자": "회계기간 외 일자 → 기간 귀속 검토, 마감(cut-off) 통제 점검, 필요 시 수정분개",
    "결산전표이상": "월할계상 이탈/중복 → 계상 근거 확인(입력 오류·자산 처분 여부), 이중 계상이면 취소분개",
    "라운드금액": "라운드 거액 → 계약서·거래명세서 대조로 실재성 확인, 특수관계자 거래 여부 검토",
    "리스의심": "리스 의심 정기 지급 → 계약서로 자산 특정·사용통제권·리스기간 확인. K-IFRS 1116 대상이면 사용권자산·리스부채 계상 및 감가상각비·이자비용 재분류, K-GAAP이면 계정 재분류 검토",
    "전표번호중복": "전표번호 중복 → 시스템 채번 오류 여부 확인",
    "역분개의심": "당일 반대전표 쌍 → 취소·재입력 사유 확인, 원거래 증빙 추적",
    "날짜이상": "날짜 형식 오류 → 원본 데이터 보정 후 재분석",
}


def add_sheet(src, df, issues_df):
    backup = None
    if src.lower().endswith((".xlsx", ".xls")):
        if glob.glob(src.rsplit(".", 1)[0] + "_backup_*.xlsx"):
            backup = None  # 기존 백업 보존 — 시트 추가된 파일로 오염 방지
        else:
            backup = src.rsplit(".", 1)[0] + f"_backup_{date.today():%Y%m%d}.xlsx"
            shutil.copyfile(src, backup)
        target = src
    else:
        target = os.path.join(os.path.dirname(src) or ".", f"이상전표_{date.today():%Y%m%d}.xlsx")
    grouped = issues_df.groupby("idx").agg({
        "이상유형": lambda s: ", ".join(dict.fromkeys(s)),
        "이상징후": lambda s: ", ".join(dict.fromkeys(s)),
        "문제설명": lambda s: " | ".join(dict.fromkeys(s))})
    sheet = df.loc[grouped.index].drop(columns=["연월", "요일", "공휴일명", "적요s", "금액"], errors="ignore").copy()
    sheet[["이상유형", "이상징후", "문제설명"]] = grouped[["이상유형", "이상징후", "문제설명"]].values
    sheet = sheet.sort_values("전표행번호")
    mode = "a" if target == src else "w"
    kw = {"if_sheet_exists": "replace"} if mode == "a" else {}
    with pd.ExcelWriter(target, engine="openpyxl", mode=mode, **kw) as w:
        sheet.to_excel(w, sheet_name="이상전표", index=False)
    return sheet, target, backup


def img64(p):
    with open(p, "rb") as f:
        return base64.b64encode(f.read()).decode()


def write_reports(outdir, src_name, fy, df, issues_df, charts, piv, cf_chart, loan_rows,
                  cf_interp, ratio_rows, months, policy_notes):
    today = f"{date.today():%Y-%m-%d}"
    title = f"분개장 분석 보고서 — {src_name}"
    note = "본 보고서의 결과는 이상 전표의 '탐지'이며 확정 판정이 아닙니다. 최종 판단은 담당자 소명과 증빙 확인 후 이루어져야 합니다."
    summary = [(k, int(v)) for k, v in issues_df["이상유형"].value_counts().items()]
    recommends = list(dict.fromkeys(RECOMMEND_MAP[k] for k, _ in summary if k in RECOMMEND_MAP))
    grouped_sheet = issues_df.groupby("idx").agg({"이상징후": lambda s: ", ".join(dict.fromkeys(s))})
    top = df.loc[grouped_sheet.index].assign(이상징후=grouped_sheet["이상징후"].values)
    top = top.sort_values("금액", ascending=False).drop_duplicates("회계전표번호").head(15)
    cf_rows = [(m, *(piv.loc[m, c] / 1e8 for c in ("영업활동", "투자활동", "재무활동", "순증감", "누적")))
               for m in months]

    # ---- Markdown ----
    L = [f"# {title} ({today})", "", f"> {note}", "", "## 1. 요약", "",
         f"- 대상: {fy} 회계연도 분개장 (전표 {df['회계전표번호'].nunique():,}매 / {len(df):,}행)",
         f"- 이상 플래그 {len(issues_df):,}건 / 이상 전표행 {issues_df['idx'].nunique():,}건 → 원본 `이상전표` 시트 참조", "",
         "| 이상유형 | 이상징후 | 건수 |", "|---|---|---|"]
    L += [f"| {k} | {SHORT[k]} | {v} |" for k, v in summary]
    L += ["", "## 2. 주요 이상 전표 (금액 상위 15)", "",
          "| 전표번호 | 회계일 | 계정과목 | 금액 | 이상징후 |", "|---|---|---|---|---|"]
    L += [f"| {r['회계전표번호']} | {pd.Timestamp(r['회계일']).date()} | {r['계정과목']} | {r['금액']:,.0f} | {r['이상징후']} |"
          for _, r in top.iterrows()]
    L += ["", "## 3. 계정과목별 월별 추세와 해석", ""]
    if policy_notes:
        L += ["**월할 계상 정책 인식:** " + " / ".join(policy_notes), ""]
    for acc, fn, interp in charts:
        L += [f"### {acc}", f"![{acc}](charts/{os.path.basename(fn)})", "", f"- {interp}", ""]
    L += ["## 4. 현금흐름 분석", "", "![현금흐름](charts/현금흐름_월별.png)", "",
          "| 월 | 영업활동 | 투자활동 | 재무활동 | 순증감 | 누적 |", "|---|---|---|---|---|---|"]
    L += [f"| {m[5:]}월 | {o:+,.1f} | {i:+,.1f} | {f:+,.1f} | {n:+,.1f} | {c:+,.1f} |"
          for m, o, i, f, n, c in cf_rows]
    L += ["", "(단위: 억원, 현금·예금 계정 기준)", ""]
    if loan_rows:
        L += ["**차입금 변동 내역**", "", "| 일자 | 구분 | 금액(원) | 적요 | 비고 |", "|---|---|---|---|---|"]
        L += [f"| {d} | {k} | {a} | {mm} | {n} |" for d, k, a, mm, n in loan_rows]
        L += [""]
    L += [f"- {t}" for t in cf_interp]
    L += ["", "## 5. 재무비율 (기말 순증감 기준 — 기초잔액 미포함)", "", "| 비율 | 값 | 해석 |", "|---|---|---|"]
    L += [f"| {a} | {b} | {c} |" for a, b, c in ratio_rows]
    L += ["", "## 6. 권고 절차", ""] + [f"- {t}" for t in recommends]
    md_path = os.path.join(outdir, f"분석보고서_{today.replace('-', '')}.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(L) + "\n")

    # ---- Word ----
    docx_path = None
    try:
        from docx import Document
        from docx.shared import Cm, Pt
        from docx.oxml.ns import qn
        doc = Document()
        st = doc.styles["Normal"]; st.font.name = "맑은 고딕"; st.font.size = Pt(10)
        st.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        for h in ("Title", "Heading 1", "Heading 2", "Heading 3"):
            s = doc.styles[h]; s.font.name = "맑은 고딕"
            s.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "맑은 고딕")

        def tbl(header, rows):
            t = doc.add_table(rows=1, cols=len(header)); t.style = "Light Grid Accent 1"
            for i, h in enumerate(header):
                t.rows[0].cells[i].text = h
            for row in rows:
                cells = t.add_row().cells
                for i, v in enumerate(row):
                    cells[i].text = str(v)
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(8.5)

        doc.add_heading(title, 0)
        doc.add_paragraph(f"작성일: {today} / 대상: {fy} 회계연도 (전표 {df['회계전표번호'].nunique():,}매, {len(df):,}행)")
        doc.add_paragraph(note).italic = True
        doc.add_heading("1. 요약", 1)
        tbl(["이상유형", "이상징후", "건수"], [(k, SHORT[k], v) for k, v in summary])
        doc.add_heading("2. 주요 이상 전표 (금액 상위 15)", 1)
        tbl(["전표번호", "회계일", "계정과목", "금액(원)", "이상징후"],
            [(r["회계전표번호"], pd.Timestamp(r["회계일"]).date(), r["계정과목"], f"{r['금액']:,.0f}", r["이상징후"])
             for _, r in top.iterrows()])
        doc.add_heading("3. 계정과목별 월별 추세와 해석", 1)
        for acc, fn, interp in charts:
            doc.add_heading(acc, 3)
            doc.add_picture(fn, width=Cm(15.5))
            doc.add_paragraph(interp).runs[0].font.size = Pt(9)
        doc.add_heading("4. 현금흐름 분석", 1)
        doc.add_picture(cf_chart, width=Cm(16.5))
        tbl(["월", "영업활동", "투자활동", "재무활동", "순증감", "누적"],
            [(m[5:] + "월", f"{o:+,.1f}", f"{i:+,.1f}", f"{f:+,.1f}", f"{n:+,.1f}", f"{c:+,.1f}")
             for m, o, i, f, n, c in cf_rows])
        doc.add_paragraph("(단위: 억원, 현금·예금 계정 기준)").runs[0].font.size = Pt(8)
        if loan_rows:
            doc.add_heading("차입금 변동 내역", 3)
            tbl(["일자", "구분", "금액(원)", "적요", "비고"], loan_rows)
        for t in cf_interp:
            doc.add_paragraph(t, style="List Bullet").runs[0].font.size = Pt(9)
        doc.add_heading("5. 재무비율 (기말 순증감 기준 — 기초잔액 미포함)", 1)
        tbl(["비율", "값", "해석"], ratio_rows)
        doc.add_heading("6. 권고 절차", 1)
        for t in recommends:
            doc.add_paragraph(t, style="List Bullet")
        docx_path = os.path.join(outdir, f"분석보고서_{today.replace('-', '')}.docx")
        doc.save(docx_path)
    except ImportError:
        print("[경고] python-docx 없음 — Word 생략(설치 권장: pip install python-docx)")

    # ---- HTML → PDF ----
    H = ["""<meta charset='utf-8'><style>
body{font-family:'Malgun Gothic';font-size:10.5pt;margin:28px;color:#222}
h1{font-size:17pt;border-bottom:2px solid #345;padding-bottom:6px}
h2{font-size:13pt;color:#234;margin-top:22px} h3{font-size:11pt;margin:14px 0 4px}
table{border-collapse:collapse;width:100%;font-size:9pt;margin:8px 0}
th,td{border:1px solid #999;padding:4px 6px;text-align:left}th{background:#eef}
img{width:100%;max-width:640px;display:block;margin:4px 0}
.note{color:#666;font-style:italic}.interp{font-size:9.5pt;color:#333;margin:2px 0 12px}
.chart{page-break-inside:avoid}</style>"""]
    H.append(f"<h1>{title}</h1><p>작성일: {today} / 대상: {fy} 회계연도 "
             f"(전표 {df['회계전표번호'].nunique():,}매, {len(df):,}행)</p><p class='note'>{note}</p>")
    H.append(f"<h2>1. 요약</h2><p>이상 플래그 {len(issues_df):,}건, 이상 전표행 {issues_df['idx'].nunique():,}건 탐지. "
             "상세는 원본 '이상전표' 시트 참조.</p>")
    H.append("<table><tr><th>이상유형</th><th>이상징후</th><th>건수</th></tr>"
             + "".join(f"<tr><td>{k}</td><td>{SHORT[k]}</td><td>{v}</td></tr>" for k, v in summary) + "</table>")
    H.append("<h2>2. 주요 이상 전표 (금액 상위 15)</h2>")
    H.append("<table><tr><th>전표번호</th><th>회계일</th><th>계정과목</th><th>금액(원)</th><th>이상징후</th></tr>"
             + "".join(f"<tr><td>{r['회계전표번호']}</td><td>{pd.Timestamp(r['회계일']).date()}</td><td>{r['계정과목']}</td>"
                       f"<td>{r['금액']:,.0f}</td><td>{r['이상징후']}</td></tr>" for _, r in top.iterrows()) + "</table>")
    H.append("<h2>3. 계정과목별 월별 추세와 해석</h2>")
    for acc, fn, interp in charts:
        H.append(f"<div class='chart'><h3>{acc}</h3><img src='data:image/png;base64,{img64(fn)}'>"
                 f"<p class='interp'>▶ {interp}</p></div>")
    H.append("<h2>4. 현금흐름 분석</h2>")
    H.append(f"<div class='chart'><img src='data:image/png;base64,{img64(cf_chart)}' style='max-width:760px'></div>")
    H.append("<table><tr><th>월</th><th>영업활동</th><th>투자활동</th><th>재무활동</th><th>순증감</th><th>누적</th></tr>"
             + "".join(f"<tr><td>{m[5:]}월</td><td>{o:+,.1f}</td><td>{i:+,.1f}</td><td>{f:+,.1f}</td>"
                       f"<td>{n:+,.1f}</td><td>{c:+,.1f}</td></tr>" for m, o, i, f, n, c in cf_rows)
             + "</table><p class='note'>(단위: 억원, 현금·예금 계정 기준)</p>")
    if loan_rows:
        H.append("<h3>차입금 변동 내역</h3>")
        H.append("<table><tr><th>일자</th><th>구분</th><th>금액(원)</th><th>적요</th><th>비고</th></tr>"
                 + "".join(f"<tr><td>{d}</td><td>{k}</td><td>{a}</td><td>{mm}</td><td>{n}</td></tr>"
                           for d, k, a, mm, n in loan_rows) + "</table>")
    H.append("<ul>" + "".join(f"<li>{t}</li>" for t in cf_interp) + "</ul>")
    H.append("<h2>5. 재무비율 (기말 순증감 기준 — 기초잔액 미포함)</h2>")
    H.append("<table><tr><th>비율</th><th>값</th><th>해석</th></tr>"
             + "".join(f"<tr><td>{a}</td><td>{b}</td><td>{c}</td></tr>" for a, b, c in ratio_rows) + "</table>")
    H.append("<h2>6. 권고 절차</h2><ul>" + "".join(f"<li>{t}</li>" for t in recommends) + "</ul>")
    html_path = os.path.join(outdir, "분석보고서.html")
    with open(html_path, "w", encoding="utf-8") as f:
        f.write("\n".join(H))

    pdf_path = os.path.join(outdir, f"분석보고서_{today.replace('-', '')}.pdf")
    edge = shutil.which("msedge") or next((p for p in (
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe") if os.path.exists(p)), None)
    if edge:
        subprocess.run([edge, "--headless", "--disable-gpu", "--no-pdf-header-footer",
                        f"--print-to-pdf={pdf_path}", "file:///" + html_path.replace(os.sep, "/")],
                       capture_output=True, timeout=120)
        if not os.path.exists(pdf_path):
            pdf_path = None
    else:
        pdf_path = None
        print("[경고] Edge 미발견 — PDF 생략. 수동 변환: 브라우저에서 분석보고서.html 열어 PDF로 인쇄")
    return md_path, docx_path, pdf_path


def main():
    ap = argparse.ArgumentParser(description="분개장 분석 파이프라인")
    ap.add_argument("--input", required=True, help="분개장 파일 경로 (xlsx/xls/csv)")
    ap.add_argument("--outdir", default=None, help="출력 폴더 (기본: <입력파일명>_분석)")
    ap.add_argument("--sheet", default=None, help="엑셀 시트명 (기본: '분개'/'전표' 포함 시트 또는 첫 시트)")
    ap.add_argument("--fy", type=int, default=None, help="회계연도 (기본: 데이터 최빈 연도)")
    args = ap.parse_args()

    src = os.path.abspath(args.input)
    outdir = args.outdir or os.path.join(os.path.dirname(src),
                                         os.path.splitext(os.path.basename(src))[0] + "_분석")
    chart_dir = os.path.join(outdir, "charts")
    os.makedirs(chart_dir, exist_ok=True)

    df, fy, _sheet = load(src, args.sheet, args.fy)
    print(f"[기간] {df['회계일'].min().date()} ~ {df['회계일'].max().date()} / "
          f"전표 {df['회계전표번호'].nunique():,}매, {len(df):,}행")

    issues_df, policy_notes = detect(df, fy)
    print(f"[탐지] 플래그 {len(issues_df)}건, 이상 행 {issues_df['idx'].nunique()}건, "
          f"이상 전표 {df.loc[issues_df['idx'].unique(), '회계전표번호'].nunique()}매")
    print(issues_df["이상유형"].value_counts().to_string())
    for n in policy_notes:
        print("[정책]", n)

    charts, months = build_charts(df, issues_df, fy, chart_dir)
    print(f"[차트] 계정과목별 {len(charts)}건")
    piv, cf_chart, loan_rows, cf_interp = cashflow(df, fy, months, chart_dir)
    print(f"[현금흐름] 영업 {piv['영업활동'].sum() / 1e8:+,.1f}억 / 투자 {piv['투자활동'].sum() / 1e8:+,.1f}억 / "
          f"재무 {piv['재무활동'].sum() / 1e8:+,.1f}억 / 차입금 이벤트 {len(loan_rows)}건")
    ratio_rows = ratios(df, fy)

    m_path, m_counts = m_candidates(df, issues_df, outdir)
    print(f"[AI정황후보] {m_path} — " + ", ".join(f"{k.split('.')[0]}:{v}" for k, v in m_counts.items()))

    sheet_df, target, backup = add_sheet(src, df, issues_df)
    print(f"[시트] '이상전표' {len(sheet_df)}행 → {target}" + (f" (백업: {backup})" if backup else ""))
    md, docx, pdf = write_reports(outdir, os.path.basename(src), fy, df, issues_df, charts,
                                  piv, cf_chart, loan_rows, cf_interp, ratio_rows, months, policy_notes)
    print(f"[보고서] {md}")
    if docx:
        print(f"[보고서] {docx}")
    if pdf:
        print(f"[보고서] {pdf}")
    print("[완료] 다음 단계(AI 수행): ① AI정황_후보.md를 판독해 정황 판단(원시 데이터를 직접 읽지 말 것) "
          "② 보고서 해석문 검토·보강 ③ 권고 절차 구체화")


if __name__ == "__main__":
    sys.exit(main())
